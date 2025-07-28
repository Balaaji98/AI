from docx import Document

def create_sample_requirement_doc(file_path):
    """
    Creates a sample Word document with a requirement for scanning a new container ID number
    while picking stock using the HHT.
    
    Args:
        file_path (str): Path to save the Word document.
    """
    doc = Document()
    doc.add_heading('Requirement Document', level=1)

    doc.add_heading('Title', level=2)
    doc.add_paragraph('Add functionality to scan a new container ID number while picking stock using the HHT.')

    doc.add_heading('Background', level=2)
    doc.add_paragraph(
        'Currently, the system allows scanning of item barcodes and location IDs during the picking process. '
        'However, there is no provision to scan a container ID number. This enhancement is required to improve '
        'tracking and accuracy during the picking process.'
    )

    doc.add_heading('Requirement Details', level=2)
    doc.add_paragraph(
        '1. Modify the HHT picking workflow to include a step for scanning the container ID number.\n'
        '2. Validate the scanned container ID against the database to ensure it is valid and active.\n'
        '3. Update the RDT rules to handle the new container ID validation logic.\n'
        '4. Log the scanned container ID in the picking transaction table for audit purposes.'
    )

    doc.add_heading('Acceptance Criteria', level=2)
    doc.add_paragraph(
        '1. The HHT should prompt the user to scan a container ID after scanning the item barcode.\n'
        '2. The system should validate the container ID and display an error message if it is invalid.\n'
        '3. The container ID should be logged in the database along with the picking transaction.'
    )

    doc.add_heading('References', level=2)
    doc.add_paragraph('Refer to the WMS Dispatcher 2019 User Guide for existing picking workflows.')

    # Save the document
    doc.save(file_path)
    print(f"Sample requirement document created at: {file_path}")

# Specify the file path where the document will be saved
file_path = '/Users/balaajiganesh/Library/Mobile Documents/com~apple~CloudDocs/WMS-AI-DESIGN-AGENT/sample_requirement.docx'
create_sample_requirement_doc(file_path)